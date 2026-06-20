# ruff: noqa: E501

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Отчет ai-pricing-service.docx"


def set_paragraph_text(paragraph, text: str, *, body: bool = True) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if body:
        run.font.size = Pt(14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
        paragraph.paragraph_format.line_spacing = 1.5


def set_cell_text(cell, text: str) -> None:
    for extra_paragraph in cell.paragraphs[1:]:
        extra_paragraph._element.getparent().remove(extra_paragraph._element)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_code_cell(cell, text: str) -> None:
    for extra_paragraph in cell.paragraphs[1:]:
        extra_paragraph._element.getparent().remove(extra_paragraph._element)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0


def main() -> None:
    if not REPORT.exists():
        raise FileNotFoundError("Run the first- and second-third report scripts before the final script.")

    document = Document(REPORT)
    paragraphs = document.paragraphs
    if not paragraphs[270].text.strip().startswith("4 "):
        raise RuntimeError("Unexpected report structure at chapter 4")
    if paragraphs[344].text.strip() != "ЗАКЛЮЧЕНИЕ":
        raise RuntimeError("Unexpected report structure at conclusion")

    content = {
        272: "Наблюдаемость AI Pricing Service построена вокруг трёх источников эксплуатационной информации: структурированных журналов, метрик Prometheus и health checks. Их совместное использование позволяет связать отдельный HTTP-запрос с результатом инференса и состоянием модельных артефактов.",
        273: "Журналирование реализовано библиотекой structlog. События выводятся в JSON и содержат временную метку UTC, уровень, название события и его контекст. Машиночитаемый формат упрощает фильтрацию записей и последующую передачу в системы агрегации логов.",
        274: "Приложение регистрирует следующие основные события:",
        275: "Запуск и штатное завершение сервиса с указанием версии и окружения.",
        276: "Успешную загрузку модели либо ошибку загрузки артефактов.",
        277: "Завершение HTTP-запроса с методом, endpoint, статусом и длительностью.",
        278: "Ошибки Pydantic-валидации входных данных.",
        279: "Успешный прогноз с product_id и confidence_score либо ошибку инференса.",
        280: "Обработанные ServiceError с безопасным кодом и сообщением.",
        281: "Служебные сведения readiness-проверки и версию используемой модели.",
        282: "Middleware назначает каждому запросу request_id. Если клиент передал X-Request-ID, значение сохраняется; иначе создаётся UUID. Идентификатор добавляется в контекст structlog, тело успешного прогноза и заголовок HTTP-ответа.",
        283: "Пример фактических событий после обращения к контейнеру представлен в листинге 3.",
        284: "Листинг 3 — Пример JSON-журналирования прогноза",
        285: "Метрики формируются библиотекой prometheus-client и публикуются через GET /metrics в стандартном текстовом формате Prometheus. Отдельно учитываются характеристики HTTP-слоя и модельного инференса.",
        286: "В приложении зарегистрированы следующие показатели:",
        287: "http_requests_total — число обработанных HTTP-запросов по endpoint и статусу.",
        288: "http_request_duration_seconds — распределение длительности HTTP-запросов.",
        289: "prediction_latency_seconds — длительность непосредственно модельного прогноза.",
        290: "prediction_confidence_score — histogram значений уверенности прогнозов.",
        291: "model_prediction_confidence — последнее рассчитанное значение confidence.",
        292: "low_confidence_predictions_total — число прогнозов ниже настроенного порога.",
        293: "model_errors_total — число ошибок загрузки и инференса по типу ошибки.",
        294: "Эти метрики позволяют контролировать частоту ошибок, задержки и изменение распределения инженерной уверенности. Конфигурация prometheus/prometheus.yml опрашивает приложение по адресу ai-pricing-service:8000/metrics.",
        295: "Для разделения доступности процесса и готовности ML-контура реализованы три эксплуатационных endpoint.",
        296: "GET /health возвращает статус ok, версию сервиса и версию API. Проверка не выполняет модельный прогноз и подходит для общей диагностики HTTP-приложения.",
        297: "GET /health/live возвращает статус alive и подтверждает, что процесс способен обрабатывать запросы.",
        298: "GET /health/ready проверяет признак загрузки модели, выполняет тестовый прогноз и возвращает версию и metadata. При неготовой модели используется статус 503 и код model_not_ready.",
        299: "Dockerfile и docker-compose.yml используют /health/ready в healthcheck. Поэтому зависимый контейнер Prometheus запускается только после подтверждения работоспособности полного модельного контура.",
        300: "Реализованная наблюдаемость обеспечивает трассировку запросов, измерение задержек, контроль ошибок и автоматическую проверку готовности. Содержимое HTTP-запросов и модельные артефакты в журналы не записываются.",
        302: "Конфигурация вынесена из программной логики и разделена на группы app, api и ml. Базовые значения хранятся в configs/config.yaml, а параметры конкретного окружения могут переопределяться переменными среды.",
        303: "Классы AppConfig, ApiConfig, MlConfig и Settings реализованы на Pydantic и Pydantic Settings. Для вложенных параметров применяется разделитель двойное подчёркивание, например ML__MODEL_VERSION.",
        304: "При запуске python-dotenv загружает файл .env без перезаписи уже заданных переменных процесса. Затем YAML преобразуется в словарь, а значения окружения с префиксами APP__, API__ и ML__ заменяют соответствующие поля.",
        305: "Файл .env.example перечисляет поддерживаемые переменные и содержит только безопасные значения по умолчанию. Локальный .env исключён из Git и Docker build context.",
        306: "Конфигурация охватывает следующие параметры:",
        307: "Имя и версия приложения.",
        308: "Адрес и порт сервиса.",
        309: "Название эксплуатационного окружения.",
        310: "Версия HTTP API.",
        311: "Путь к сериализованной модели.",
        312: "Путь к вспомогательной статистике цен.",
        313: "Путь к metadata модели.",
        314: "Ожидаемая версия модельных артефактов.",
        315: "Порог формирования предупреждения о низкой уверенности.",
        316: "Вложенное переопределение значений через переменные окружения.",
        317: "Полный набор поддерживаемых настроек представлен в таблице 4.1.",
        318: "Таблица 4.1 — Параметры конфигурации AI Pricing Service",
        320: "Параметр ML__MIN_CONFIDENCE_THRESHOLD валидируется в диапазоне от 0 до 1. Изменение порога влияет только на формирование warning и не изменяет сам прогноз модели.",
        321: "Относительные пути к артефактам разрешаются от корня проекта, а не от текущего рабочего каталога. Это обеспечивает одинаковое поведение при локальном и контейнерном запуске.",
        322: "Docker Compose явно задаёт APP__ENV=production. Остальные значения берутся из config.yaml внутри образа либо могут быть переданы контейнеру стандартными средствами Docker.",
        323: "Пример полного набора переменных окружения приведён в листинге 4.",
        324: "Листинг 4 — Пример конфигурации приложения",
        325: "Централизованная конфигурация исключает дублирование путей и версий в коде, позволяет валидировать настройки при старте и отделяет параметры развёртывания от реализации API и модели.",
        327: "Проект не обращается к внешнему AI-провайдеру и не требует API-ключей. Основные аспекты безопасности связаны с обработкой входных данных, локальными файлами окружения, модельными артефактами и запуском контейнера.",
        328: "В конфигурации отсутствуют пароли, токены и строки подключения к внешним хранилищам. Файл .env.example содержит только название сервиса, версии, сетевые параметры, пути к локальным артефактам и порог confidence.",
        329: "Механизм .env сохранён для локальных переопределений, однако реальные значения этого файла не должны публиковаться. Загруженные переменные не выводятся в журнал приложения.",
        330: "Файл .env.example позволяет воспроизвести структуру настроек без копирования локального .env и используется как явный перечень поддерживаемых параметров.",
        331: "Файл .gitignore исключает из контроля версий следующие локальные данные:",
        332: "Файл .env с переопределениями окружения.",
        333: "Виртуальные окружения Python.",
        334: "Кэши Python, Pytest, Ruff и Mypy.",
        335: "Логи, временные каталоги и промежуточные файлы.",
        336: "Каталог local_data для локальных наборов данных.",
        337: "Аналогичные правила .dockerignore не допускают попадания .env, Git metadata, виртуального окружения, кэшей и логов в контекст сборки образа.",
        338: "Структурированные логи содержат только технический контекст, необходимый для диагностики: request_id, endpoint, статус, длительность, product_id и confidence. Полное тело запроса и содержимое артефактов не журналируются.",
        339: "При загрузке вычисляются SHA-256 модели и scaler и сопоставляются со значениями metadata. Также проверяются версия, тип модели, признаки и target. Эти проверки обнаруживают повреждение или случайную подмену, но не заменяют цифровую подпись артефакта.",
        340: "Pydantic отклоняет пустой идентификатор товара, горизонт вне диапазона 1–14 и неположительную цену до передачи данных в модель. Бизнес-логика дополнительно ограничивает ценовую корректировку.",
        341: "Ошибки готовности и загрузки преобразуются в структурированный ответ без трассировки стека и внутренних путей. Контейнер запускает приложение от непривилегированного пользователя appuser.",
        342: "Проект использует синтетический обучающий набор и не хранит персональные данные. История из 50 последних прогнозов существует только в памяти процесса и очищается при перезапуске.",
        343: "Сериализация pickle требует доверенного источника: контрольная сумма подтверждает целостность относительно metadata, но загрузка стороннего pickle остаётся небезопасной. Поэтому runtime должен использовать только артефакты, поставляемые вместе с проверенным образом проекта.",
        345: "В рамках итогового проекта разработан интеллектуальный веб-сервис AI Pricing Service, который прогнозирует товарный спрос и рассчитывает рекомендованную цену через документированный HTTP-интерфейс.",
        346: "В ходе работы сформирован контракт задачи регрессии, определены входные признаки и целевая переменная, разработан воспроизводимый генератор синтетических данных и выполнен разведочный анализ 1000 наблюдений.",
        347: "Для выбора алгоритма на единой тестовой выборке сопоставлены DummyRegressor, LinearRegression, DecisionTreeRegressor и RandomForestRegressor. Все варианты использовали одинаковое разбиение 80/20 и единое преобразование категориального признака.",
        348: "Наилучший результат показал RandomForestRegressor: MAE 2,3240 и RMSE 2,9570. Модель включена в Pipeline после DictVectorizer, что обеспечивает одинаковую подготовку признаков при обучении и инференсе.",
        349: "В результате реализации проекта получены следующие практические результаты:",
        350: "Подготовлен воспроизводимый синтетический набор и выполнен EDA признаков и спроса.",
        351: "Реализован скрипт сравнения baseline и улучшенных моделей с выгрузкой метрик в JSON.",
        352: "Обучена и сохранена финальная модель RandomForestRegressor с версией v1.",
        353: "Разработаны REST API, Swagger UI и встроенный веб-интерфейс прогноза.",
        354: "Маршрут /predict использует реальный сериализованный pipeline, а не программную заглушку.",
        355: "Добавлены metadata, контрольные суммы и readiness-проверка модельных артефактов.",
        356: "Внедрены JSON-логи, X-Request-ID, Prometheus-метрики и health checks.",
        357: "Подготовлены Dockerfile и Docker Compose для запуска API вместе с Prometheus.",
        358: "Созданы автоматизированные тесты, OpenAPI-схема, model card и инструкции запуска.",
        359: "Архитектура отделяет прогноз спроса от бизнес-правила цены. Это делает происхождение результата прозрачным: ML-компонент возвращает demand_units, а recommended_price рассчитывается отдельным ограниченным алгоритмом.",
        360: "Проверка проекта включает 29 тестов, lint посредством Ruff, валидацию Docker Compose и сборку контейнерного образа. Сквозной запуск подтвердил состояние healthy и успешный ответ POST /v1/predict.",
        361: "Практическая значимость решения состоит в демонстрации полного жизненного цикла ML-сервиса: от данных и экспериментов до версионированного инференса, API, наблюдаемости и контейнерного развёртывания.",
        362: "Основное ограничение связано с синтетическими данными. Модель не учитывает реальную историю продаж, сезонность, промо-акции, остатки и конкурентные цены, а confidence_score является инженерной оценкой, а не откалиброванной вероятностью.",
        363: "Для применения к реальному ценообразованию потребуется исторический датасет, временная валидация, бизнес-метрики, мониторинг drift и механизм безопасного обновления подписанных модельных артефактов.",
        364: "Поставленная цель достигнута: создан работоспособный и воспроизводимый сервис, который использует обученную модель для прогнозирования спроса, формирует ценовую рекомендацию и предоставляет необходимые средства контроля качества и эксплуатации.",
        366: "FastAPI Documentation. — Текст : электронный // FastAPI : [сайт]. — URL: https://fastapi.tiangolo.com/ (дата обращения: 20.06.2026).",
        367: "RandomForestRegressor. — Текст : электронный // scikit-learn : [сайт]. — URL: https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.RandomForestRegressor.html (дата обращения: 20.06.2026).",
        368: "Pydantic Settings. — Текст : электронный // Pydantic : [сайт]. — URL: https://docs.pydantic.dev/latest/concepts/pydantic_settings/ (дата обращения: 20.06.2026).",
        369: "Docker Compose Documentation. — Текст : электронный // Docker Docs : [сайт]. — URL: https://docs.docker.com/compose/ (дата обращения: 20.06.2026).",
    }

    for index, text in content.items():
        body = not paragraphs[index].style.name.startswith("Heading") and paragraphs[index].style.name not in {
            "TblTitle",
            "ListCap",
        }
        set_paragraph_text(paragraphs[index], text, body=body)

    log_listing = """{"request_id":"report-final-check","product_id":"ELEC-100500","confidence_score":0.82,"event":"prediction_succeeded","level":"info","timestamp":"2026-06-19T21:29:06.586972Z"}
{"method":"POST","endpoint":"/v1/predict","status_code":200,"duration_ms":57.08,"client_ip":"172.19.0.1","event":"http_request_completed","request_id":"report-final-check","level":"info","timestamp":"2026-06-19T21:29:06.587951Z"}"""
    set_code_cell(document.tables[7].cell(0, 0), log_listing)

    config_table = document.tables[8]
    while len(config_table.rows) < 12:
        config_table.add_row()
    config_rows = [
        ("Параметр", "Переменная окружения", "Назначение"),
        ("app.name", "APP__NAME", "Название сервиса"),
        ("app.version", "APP__VERSION", "Версия приложения"),
        ("app.host", "APP__HOST", "Адрес привязки"),
        ("app.port", "APP__PORT", "Порт приложения"),
        ("app.env", "APP__ENV", "Название окружения"),
        ("api.version", "API__VERSION", "Версия API"),
        ("ml.model_path", "ML__MODEL_PATH", "Путь к модели"),
        ("ml.scaler_path", "ML__SCALER_PATH", "Путь к статистике цен"),
        ("ml.metadata_path", "ML__METADATA_PATH", "Путь к metadata"),
        ("ml.model_version", "ML__MODEL_VERSION", "Версия артефактов"),
        ("ml.min_confidence_threshold", "ML__MIN_CONFIDENCE_THRESHOLD", "Порог warning"),
    ]
    for row_index, values in enumerate(config_rows):
        for column_index, value in enumerate(values):
            set_cell_text(config_table.cell(row_index, column_index), value)

    env_listing = """APP__NAME=ai-pricing-service
APP__VERSION=1.0.0
APP__HOST=0.0.0.0
APP__PORT=8000
APP__ENV=production
API__VERSION=v1
ML__MODEL_PATH=./artifacts/model_v1.pkl
ML__SCALER_PATH=./artifacts/scaler_v1.pkl
ML__METADATA_PATH=./artifacts/model_metadata.json
ML__MODEL_VERSION=v1
ML__MIN_CONFIDENCE_THRESHOLD=0.6"""
    set_code_cell(document.tables[9].cell(0, 0), env_listing)

    checklist = document.tables[10]
    checklist_rows = [
        ("№", "Критерий", "Да/Нет", "Где смотреть"),
        ("1", "Сервис запускается по инструкции из README и работает", "+", "README.md; Dockerfile; docker-compose.yml; src/app/main.py"),
        ("2", "/predict использует реальную модель, а не заглушку", "+", "src/app/api/predict.py; src/app/api/v1/predict.py; src/app/services/model_loader.py"),
        ("3", "Есть EDA и хотя бы один эксперимент с метриками", "+", "notebooks/01_eda.ipynb; разделы 1.3 и 2.3; reports/model_comparison.json"),
        ("4", "Есть baseline и улучшенная модель, сравнение по метрикам", "+", "Разделы 2.1–2.4; таблицы 2.1 и 2.2; scripts/evaluate_models.py"),
        ("5", "Код структурирован в src/, а не свален в один ноутбук", "+", "src/app/api; src/app/core; src/app/models; src/app/services"),
        ("6", "Есть Dockerfile или внятный сценарий развёртывания", "+", "Dockerfile; docker-compose.yml; README.md, раздел запуска через Docker"),
        ("7", "Есть .env.example, нет реальных секретов в репозитории", "+", ".env.example; .gitignore; .dockerignore; src/app/core/config.py"),
        ("8", "Реализованы логи и эндпоинт /health", "+", "src/app/core/logger.py; src/app/core/metrics.py; src/app/api/health.py"),
        ("9", "Обоснован выбор финальной модели", "+", "Раздел 2.4; artifacts/model_metadata.json; reports/model_comparison.json"),
        ("10", "README и отчёт позволяют понять сценарий демонстрации", "+", "README.md; разделы 3.2–3.3; заключение отчёта"),
    ]
    for row_index, values in enumerate(checklist_rows):
        for column_index, value in enumerate(values):
            set_cell_text(checklist.cell(row_index, column_index), value)

    document.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
